"""Tests for revenue_model.docx_builder.

Smoke/contract tests in the style of test_viz.py: python-docx is optional, so
``pytest.importorskip("docx")`` keeps the suite CI-friendly when the extra is
absent. Chart-dependent cases additionally skip without matplotlib.

Assertions check the *contract* (file generates, sections present, ABC colors
applied, honest-default annotations appear) rather than pixel-level layout.
"""
import pytest

pytest.importorskip("docx")  # needs python-docx (the [docx] extra)
from docx import Document  # noqa: E402

from revenue_model.demo import build_novatech  # noqa: E402
from revenue_model.docx_builder import build_docx  # noqa: E402


@pytest.fixture
def model():
    return build_novatech()


def _fulltext(path):
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts), doc


def _all_runs(doc):
    for p in doc.paragraphs:
        yield from p.runs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield from p.runs


# ---- core generation --------------------------------------------------------
def test_build_docx_generates_nonempty_file(model, tmp_path):
    out = build_docx(model, str(tmp_path / "m.docx"))
    import os
    assert os.path.exists(out)
    assert os.path.getsize(out) > 5000  # a real doc, not a stub


def test_seven_sections_present_en(model, tmp_path):
    build_docx(model, str(tmp_path / "m.docx"), lang="en")
    ft, _ = _fulltext(tmp_path / "m.docx")
    for title in ["Executive Summary", "Company & Segment Overview",
                  "Segment Driver Tables", "Residual Alignment",
                  "Uncertainty & Scenarios", "Limitations",
                  "Methodology"]:
        assert title in ft, f"missing section: {title}"


def test_sections_present_zh(model, tmp_path):
    build_docx(model, str(tmp_path / "m.docx"), lang="zh")
    ft, _ = _fulltext(tmp_path / "m.docx")
    for title in ["执行摘要", "公司与分部概述", "分部 driver 表",
                  "差额对齐", "不确定性与情景分析", "局限性声明", "方法论"]:
        assert title in ft, f"missing zh section: {title}"


# ---- ABC coloring -----------------------------------------------------------
def test_abc_coloring_applied(model, tmp_path):
    """NovaTech has C-grade drivers (share, price) → red runs must exist."""
    build_docx(model, str(tmp_path / "m.docx"))
    _, doc = _fulltext(tmp_path / "m.docx")
    colors = set()
    for run in _all_runs(doc):
        try:
            rgb = run.font.color.rgb
            if rgb is not None:
                colors.add(str(rgb))
        except Exception:
            pass
    assert "FF0000" in colors, "C-grade red not found — ABC coloring failed"


# ---- honest defaults --------------------------------------------------------
def test_default_ranges_annotated(model, tmp_path):
    """ranges=None must surface the 'illustrative only' caveat (honest default)."""
    build_docx(model, str(tmp_path / "m.docx"))  # ranges defaults to None
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "illustrative only" in ft
    assert "default ±10%" in ft


def test_user_ranges_no_default_warning(model, tmp_path):
    build_docx(model, str(tmp_path / "m.docx"),
              ranges={"NovaTech 国内市占率": (0.10, 0.18)})
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "user-supplied" in ft
    assert "illustrative only" not in ft


def test_forecast_none_is_history_only(model, tmp_path):
    build_docx(model, str(tmp_path / "m.docx"))  # forecast_years=None
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "historical model only" in ft


def test_forecast_placeholder_when_drivers_missing(model, tmp_path):
    """forecast_years passed but drivers lack those years → prominent alarm."""
    build_docx(model, str(tmp_path / "m.docx"), forecast_years=[2025, 2026])
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "not yet populated" in ft
    assert "2025" in ft and "2026" in ft


# ---- language / charts ------------------------------------------------------
def test_lang_footnote_present(model, tmp_path):
    """The 'must remind the user' requirement: every memo carries the lang hint."""
    for lang in ["en", "zh"]:
        build_docx(model, str(tmp_path / f"m_{lang}.docx"), lang=lang)
        ft, _ = _fulltext(tmp_path / f"m_{lang}.docx")
        assert 'lang="' in ft, f"lang footnote missing in {lang}"


def test_invalid_lang_raises(model, tmp_path):
    with pytest.raises(ValueError):
        build_docx(model, str(tmp_path / "m.docx"), lang="fr")


def test_include_charts_false_table_fallback(model, tmp_path):
    """No matplotlib needed; §5 falls back to percentile tables."""
    build_docx(model, str(tmp_path / "m.docx"), include_charts=False)
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "matplotlib is not installed" in ft
    assert "P5" in ft  # percentile table rendered


def test_include_charts_true_needs_matplotlib(model, tmp_path, monkeypatch):
    """When charts requested but matplotlib absent → ImportError naming the extra.

    Uses monkeypatch so sys.modules is auto-restored — no leakage into later
    tests (a manual poison/restore corrupted matplotlib state for test_viz)."""
    import sys
    monkeypatch.setitem(sys.modules, "matplotlib", None)
    monkeypatch.setitem(sys.modules, "matplotlib.pyplot", None)
    with pytest.raises(ImportError) as excinfo:
        build_docx(model, str(tmp_path / "m.docx"), include_charts=True)
    assert "[docx]" in str(excinfo.value)


# ---- chart-embedding smoke (needs matplotlib) -------------------------------
def test_charts_embedded_when_mpl_available(model, tmp_path):
    pytest.importorskip("matplotlib")
    import matplotlib
    matplotlib.use("Agg")
    build_docx(model, str(tmp_path / "m.docx"), include_charts=True)
    # embedded images appear as inline shapes
    doc = Document(str(tmp_path / "m.docx"))
    assert len(doc.inline_shapes) >= 2, "expected ≥2 embedded charts"


# ---- hyperlinks / traceability ---------------------------------------------
def test_source_url_renders_external_hyperlinks(model, tmp_path):
    """Drivers with source_url → memo contains clickable external hyperlinks
    (the traceability that gives the memo its authority)."""
    from docx.oxml.ns import qn
    build_docx(model, str(tmp_path / "m.docx"), lang="en")
    doc = Document(str(tmp_path / "m.docx"))
    external = [h for h in doc.element.body.findall(".//" + qn("w:hyperlink"))
                if h.get(qn("r:id"))]
    assert len(external) >= 4, "expected ≥4 external hyperlinks (driver sources + GitHub docs)"


def test_source_index_appendix_present(model, tmp_path):
    build_docx(model, str(tmp_path / "m.docx"), lang="en")
    ft, _ = _fulltext(tmp_path / "m.docx")
    assert "Data Source Index" in ft
    build_docx(model, str(tmp_path / "m_zh.docx"), lang="zh")
    ft2, _ = _fulltext(tmp_path / "m_zh.docx")
    assert "数据来源索引" in ft2


def test_internal_cross_reference_links_present(model, tmp_path):
    """Executive summary carries internal links to §3/§4/§5 (chapter jumps)."""
    from docx.oxml.ns import qn
    build_docx(model, str(tmp_path / "m.docx"), lang="en")
    doc = Document(str(tmp_path / "m.docx"))
    anchors = [h.get(qn("w:anchor"))
               for h in doc.element.body.findall(".//" + qn("w:hyperlink"))
               if h.get(qn("w:anchor"))]
    assert "sec3" in anchors and "sec4" in anchors and "sec5" in anchors


def test_methodology_github_links_present(model, tmp_path):
    """Methodology section links to project docs on GitHub."""
    build_docx(model, str(tmp_path / "m.docx"), lang="en")
    doc = Document(str(tmp_path / "m.docx"))
    github_urls = [str(r.target_ref) for r in doc.part.rels.values()
                   if "hyperlink" in r.reltype and "github.com" in str(r.target_ref)]
    assert len(github_urls) >= 2, "expected ≥2 GitHub doc links in methodology"
