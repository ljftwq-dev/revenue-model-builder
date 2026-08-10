"""docx_builder — render a RevenueModel into a Word research memo.

Companion to ``excel_builder``: the .xlsx is the **working paper**
(calculation), this is the **deliverable** (narrative). Produces a 7-section
analyst memo with ABC-colored driver tables, residual alignment, embedded
uncertainty charts (Monte Carlo distribution / tornado / forecast trajectory),
and an honest Limitations section.

Optional dependency: python-docx (+ matplotlib when ``include_charts=True``).
Install via ``pip install -e ".[docx]"``. This module is **not** imported by
``import revenue_model`` — import it explicitly so the zero-dependency core
stays clean. If python-docx is missing, a clear error names the extra.

.. tip::
   ``lang`` selects the memo language: ``"en"`` (default, global/PyPI) or
   ``"zh"``. Example::

       build_docx(model, "memo.docx", lang="zh")

The generated memo carries a footnote showing the active language and how to
switch.
"""

# Lazy import python-docx — keeps ``import revenue_model`` zero-dependency.
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE
except ImportError as exc:  # pragma: no cover - exercised only without python-docx
    raise ImportError(
        "revenue_model.docx_builder requires python-docx. "
        'Install it with:  pip install -e ".[docx]"'
    ) from exc

from io import BytesIO
from typing import Dict, List, Literal, Optional, Sequence

from .driver import LEVEL_A, LEVEL_B, LEVEL_C, Driver
from .model import RevenueModel, YearResult
from .segment import Segment

_YI = 100.0  # million yuan -> 亿元 (zh) / RMB 100mn (en); matches viz.py

# ABC colors — identical to excel_builder so a reviewer reads both formats
# with the same visual convention (A black, B blue, C red).
_COLOR_A = RGBColor(0x00, 0x00, 0x00)
_COLOR_B = RGBColor(0x00, 0x00, 0xFF)
_COLOR_C = RGBColor(0xFF, 0x00, 0x00)
_LEVEL_COLOR = {LEVEL_A: _COLOR_A, LEVEL_B: _COLOR_B, LEVEL_C: _COLOR_C}

_NAVY = RGBColor(0x1F, 0x3A, 0x5F)      # section headings (excel_builder title color)
_SEG_BLUE = RGBColor(0x2E, 0x5C, 0x8A)  # segment sub-headers
_WARN_AMBER = RGBColor(0xB8, 0x86, 0x00)
_PLACEHOLDER_RED = RGBColor(0xC0, 0x00, 0x00)
_GREY = RGBColor(0x80, 0x80, 0x80)
_LINK_BLUE = "0000FF"

_FONT = "Microsoft YaHei"  # 中英文 + 东亚字体统一（硬规则）

# Project docs on GitHub — referenced from the methodology section so each
# principle is one click away from its full treatment.
_GITHUB_BASE = "https://github.com/ljftwq-dev/revenue-model-builder"
_DOC_URLS = {
    "design_principles": _GITHUB_BASE + "/blob/main/docs/design-principles.md",
    "industry_fit": _GITHUB_BASE + "/blob/main/docs/industry-fit-analysis.md",
    "extractor_proposal": _GITHUB_BASE + "/blob/main/docs/proposal-segment-extraction.md",
    "changelog": _GITHUB_BASE + "/blob/main/CHANGELOG.md",
}

_bookmark_counter = [0]


def _hyperlink_run(text, *, size=10, color=_LINK_BLUE, underline=True):
    """The <w:r> inside a hyperlink — blue underlined YaHei, half-point size."""
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(attr), _FONT)
    rPr.append(rFonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    return run


def _add_hyperlink(paragraph, url, text, *, size=10, color=_LINK_BLUE):
    """A real clickable external hyperlink (blue, underlined) via oxml.

    python-docx ships no native hyperlink API, so we build the <w:hyperlink>
    element and the external relationship directly. ``text`` shows as a blue
    underlined link; clicking opens ``url`` in the browser.
    """
    r_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.append(_hyperlink_run(text, size=size, color=color))
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_internal_link(paragraph, anchor, text, *, size=10.5, color="1F3A5F"):
    """A clickable link to an internal bookmark (<w:hyperlink w:anchor=...>).

    Used for chapter cross-references (e.g. executive summary → "see §3"). The
    target bookmark is set on a heading via the ``bookmark=`` kwarg."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.append(_hyperlink_run(text, size=size, color=color, underline=False))
    paragraph._p.append(hyperlink)
    return hyperlink


def _bookmark(paragraph, name):
    """Wrap a paragraph in <w:bookmarkStart>/<w:bookmarkEnd> for internal links."""
    _bookmark_counter[0] += 1
    bid = str(_bookmark_counter[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


Lang = Literal["zh", "en"]

# ---- bilingual strings ------------------------------------------------------
_STRINGS = {
    "zh": {
        "title_suffix": "收入模型研究 memo",
        "lang_footnote": '语言: zh（设 lang="en" 切换英文版）',
        "generated": "生成于",
        # section titles
        "sec1": "一、执行摘要",
        "sec2": "二、公司与分部概述",
        "sec3": "三、分部 driver 表",
        "sec4": "四、差额对齐",
        "sec5": "五、不确定性与情景分析",
        "sec6": "六、局限性声明",
        "sec7": "附录：方法论",
        # exec summary
        "company": "公司",
        "years_modeled": "建模年份",
        "total_reported": "总收入(年报，亿元)",
        "seg_count": "分部数",
        "residual_ratio": "差额比(最新年)",
        "scenarios_head": "情景(最新年，亿元)",
        "no_warnings": "对齐通过，无告警",
        "warnings_head": "告警",
        # overview
        "overview_placeholder": "[待填：公司分部概述 — 通过 company_overview 参数传入]",
        "segment": "分部",
        "rev_share": "收入占比(最新年)",
        # driver table
        "driver": "driver",
        "kind": "类型",
        "unit": "单位",
        "grade": "等级",
        "source": "来源",
        "seg_revenue": "分部收入",
        # residual
        "year": "年份",
        "seg_sum": "Σ分项(亿元)",
        "residual": "差额(亿元)",
        # uncertainty
        "ranges_user": "区间来源：用户传入（反映 ABC 分级）",
        "ranges_default": "⚠ 区间来源：默认 ±10%（仅示意，非真实不确定性）。真实分析请传入反映 ABC 分级的 ranges。",
        "dist_subtitle": "收入分布（蒙特卡洛）",
        "tornado_subtitle": "敏感度（龙卷风，单 driver 摆动）",
        "forecast_subtitle": "收入轨迹",
        "bear": "悲观", "base": "基准", "bull": "乐观",
        "forecast_none": "未指定预测年份 —— 本 memo 仅记录历史模型。driver 外推后传入 forecast_years=[...] 可扩展预测节。",
        "forecast_placeholder": "[预测 driver 未填充 —— 请先对 {years} 运行 extrapolate_* 再重新生成 memo]",
        "percentile": "百分位",
        "revenue_yi": "收入(亿元)",
        # limitations
        "lim_intro": "本 memo 由 revenue-model-builder 自动生成。下列局限务必在引用结论前知悉：",
        "lim_c_grade": "下列 driver 为 C 级估算（红色），不确定性最高，是预测的主要风险源：",
        "lim_none_c": "本模型无 C 级 driver。",
        "lim_default_ranges": "蒙特卡洛区间采用默认 ±10% 区间，未反映各 driver 的真实 ABC 分级不确定性 —— 分布形状结构有效，绝对宽度未校准。",
        "lim_user_ranges": "蒙特卡洛区间由用户传入，应已反映各 driver 的 ABC 分级。",
        "lim_no_charts": "本 memo 因未安装 matplotlib 而以表格形式呈现不确定性节（include_charts=False）。",
        "lim_residual_high": "最新年差额比为 {ratio:.0%}，高于 30% 健康上限 —— 未建模业务占比较大，考虑补充分部。",
        "lim_residual_ok": "差额比处于健康区间（10–30%）。",
        "lim_diagnostic_en": "程序生成的对齐告警保留英文（技术诊断信息）。",
        # methodology
        "method_body": [
            "收入分解公式：分部收入 = 市场基数 × 渗透率 × 市占率 × 单价；总收入 = Σ分部 + 差额行。",
            "ABC 数据等级：A=年报硬数据(黑)、B=第三方行业数据(蓝)、C=估算(红)。每个数字标注来源，让模型可追溯、可审计。",
            "渗透率/市占率预测用增量法（+X 个百分点/年），不用增速法 —— 有界变量用增速会指数爆炸。详见 design-principles.md 原则 3。",
            "差额行是结构性设计，绝不反推：用行业真实渗透率，模型分项必然 ≠ 年报总收入，差额吸收未建模业务。反推渗透率看似对齐，实则把失真带进预测期。",
            "先历史后预测：历史列必须先对齐，预测列保留结构直到历史 tie out。",
        ],
    },
    "en": {
        "title_suffix": "Revenue Model Research Memo",
        "lang_footnote": 'Language: en (set lang="zh" for 中文版)',
        "generated": "Generated",
        "sec1": "1. Executive Summary",
        "sec2": "2. Company & Segment Overview",
        "sec3": "3. Segment Driver Tables",
        "sec4": "4. Residual Alignment",
        "sec5": "5. Uncertainty & Scenarios",
        "sec6": "6. Limitations & Caveats",
        "sec7": "Appendix: Methodology",
        "company": "Company",
        "years_modeled": "Years modeled",
        "total_reported": "Total revenue (reported, RMB 100mn)",
        "seg_count": "Segments",
        "residual_ratio": "Residual ratio (latest yr)",
        "scenarios_head": "Scenarios (latest yr, RMB 100mn)",
        "no_warnings": "Aligned, no warnings",
        "warnings_head": "Warnings",
        "overview_placeholder": "[TODO: company & segment overview — pass via company_overview]",
        "segment": "Segment",
        "rev_share": "Revenue share (latest yr)",
        "driver": "Driver",
        "kind": "Kind",
        "unit": "Unit",
        "grade": "Grade",
        "source": "Source",
        "seg_revenue": "Segment revenue",
        "year": "Year",
        "seg_sum": "Σ segments (RMB 100mn)",
        "residual": "Residual (RMB 100mn)",
        "ranges_user": "Ranges source: user-supplied (A/B/C-reflective)",
        "ranges_default": "⚠ Ranges source: default ±10% bands (illustrative only, not real uncertainty). Pass ranges= with A/B/C-reflective intervals for real analysis.",
        "dist_subtitle": "Revenue distribution (Monte Carlo)",
        "tornado_subtitle": "Sensitivity (tornado, one-driver swing)",
        "forecast_subtitle": "Revenue trajectory",
        "bear": "Bear", "base": "Base", "bull": "Bull",
        "forecast_none": "No forecast horizon specified — this memo documents the historical model only. Add forecast_years=[...] once drivers are extrapolated to extend.",
        "forecast_placeholder": "[Forecast drivers not yet populated — run extrapolate_* for {years} before regenerating memo]",
        "percentile": "Percentile",
        "revenue_yi": "Revenue (RMB 100mn)",
        "lim_intro": "This memo was auto-generated by revenue-model-builder. Read the limitations below before citing any conclusion:",
        "lim_c_grade": "The following drivers are C-grade estimates (red) — highest uncertainty, the main forecast risk:",
        "lim_none_c": "No C-grade drivers in this model.",
        "lim_default_ranges": "Monte Carlo intervals used the default ±10% bands, not driver-specific uncertainty — the distribution shape is structurally valid but absolute spread is not calibrated to data grades.",
        "lim_user_ranges": "Monte Carlo intervals were user-supplied; they should already reflect each driver's A/B/C grade.",
        "lim_no_charts": "This memo renders the uncertainty section as tables because matplotlib is not installed (include_charts=False).",
        "lim_residual_high": "The latest-year residual ratio is {ratio:.0%}, above the 30% healthy ceiling — unmodeled business dominates; consider adding segments.",
        "lim_residual_ok": "Residual ratio is in the healthy 10–30% band.",
        "lim_diagnostic_en": "Program-generated alignment warnings are shown verbatim (technical diagnostics).",
        "method_body": [
            "Decomposition: segment_revenue = base × penetration × share × price; total_revenue = Σ(segments) + residual.",
            "ABC data grading: A=annual-report hard data (black), B=third-party industry data (blue), C=estimate (red). Every number carries a source — the model is auditable.",
            "Forecast bounded ratios (penetration/share) in absolute increments (+pp/yr), never growth rates — a growth rate on a bounded variable explodes past 100%. See design-principles.md Principle 3.",
            "The residual line is structural, never back-solved: with industry-realistic penetration, modeled segments never exactly equal reported total. The residual absorbs unmodeled business. Back-solving penetration to force alignment poisons the forecast period.",
            "History first, then forecast: historical columns must tie out before forecast columns are filled.",
        ],
    },
}


def _check_lang(lang: str) -> None:
    if lang not in _STRINGS:
        raise ValueError(f"lang must be 'zh' or 'en', got {lang!r}")


# ---- low-level formatting helpers ------------------------------------------
def _style_run(run, *, size=10.5, bold=False, italic=False, color=None,
               font: str = _FONT):
    """Apply font + east-asian glyph fallback to a run.

    Microsoft YaHei is set as *both* the western and east-asian font, so
    Chinese glyphs render instead of falling back to a default that drops CJK.
    """
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def _heading(doc, text, *, level=1, bookmark=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _style_run(run, size=15 if level == 1 else 12.5, bold=True,
               color=_NAVY if level == 1 else _SEG_BLUE)
    if bookmark:
        _bookmark(p, bookmark)
    return p


def _para(doc, text, *, size=10.5, bold=False, italic=False, color=None,
          align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _bullet(doc, text, *, color=None, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        _style_run(r, size=10.5, bold=True, color=color)
    run = p.add_run(text)
    _style_run(run, size=10.5, color=color)
    return p


def _table(doc, header: List[str], rows: List[List[str]],
           row_colors: Optional[List[Optional[RGBColor]]] = None):
    """A bordered table with a bold header row. ``row_colors`` optionally sets
    the font color for every cell in a body row (used for ABC grading)."""
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for ci, h in enumerate(header):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _style_run(run, size=10, bold=True, color=_NAVY)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        color = row_colors[ri] if row_colors else None
        for ci, val in enumerate(row):
            cell = t.rows[1 + ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _style_run(run, size=10, color=color)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    return t


# ---- chart / forecast helpers ----------------------------------------------
def _import_mpl():
    """Lazy matplotlib import — only triggered when embedding charts."""
    try:
        import matplotlib
        matplotlib.use("Agg")  # headless, CI-safe (matches test_viz.py)
        import matplotlib.pyplot as plt
        return plt
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "docx_builder chart embedding requires matplotlib. "
            'Install the [docx] extra:  pip install -e ".[docx]"'
        ) from exc


def _embed_fig(doc, fig, plt, *, width_in=6.2):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_in))
    plt.close(fig)


def _default_ranges(seg: Segment, year: int) -> Dict[str, tuple]:
    """±10% bands per driver — the honest-default fallback (see design §3).
    Identical to viz._default_ranges; duplicated here so docx_builder does not
    reach into a private name."""
    out: Dict[str, tuple] = {}
    for d in seg.drivers():
        v = d.get(year)
        out[d.name] = (v * 0.9, v * 1.1)
    return out


def _model_default_ranges(model: RevenueModel, year: int) -> Dict[str, tuple]:
    out: Dict[str, tuple] = {}
    for seg in model.segments:
        out.update(_default_ranges(seg, year))
    return out


def _has_forecast_values(model: RevenueModel,
                         forecast_years: Optional[Sequence[int]]) -> bool:
    if not forecast_years:
        return False
    for seg in model.segments:
        for d in seg.drivers():
            for y in forecast_years:
                if y not in d.values:
                    return False
    return True


def _yi(v: float) -> str:
    """million yuan -> 亿元 / RMB 100mn, formatted."""
    return f"{v / _YI:.2f}"


# ---- per-section renderers -------------------------------------------------
def _add_cover(doc: Document, model: RevenueModel, s: dict, lang: str):
    import datetime
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{model.company} · {s['title_suffix']}")
    _style_run(run, size=20, bold=True, color=_NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{s['generated']} {datetime.date.today().isoformat()}")
    _style_run(r2, size=10, color=_GREY)
    # language footnote (the "must remind the user" requirement)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(s["lang_footnote"])
    _style_run(r3, size=9, italic=True, color=_GREY)


def _add_executive_summary(doc, model: RevenueModel, s: dict, lang: str,
                           results: List[YearResult]):
    _heading(doc, s["sec1"], level=1, bookmark="sec1")
    latest = results[-1]
    total_vals = [f"{model.total_revenue[y] / _YI:.2f}" for y in model.years()]
    rows = [
        [s["company"], model.company],
        [s["years_modeled"], " / ".join(str(y) for y in model.years())],
        [s["total_reported"], " / ".join(total_vals)],
        [s["seg_count"], str(len(model.segments))],
        [s["residual_ratio"], f"{latest.residual_ratio:.1%}"],
    ]
    _table(doc, ["", ""], rows)
    # warnings
    all_warnings = [w for r in results for w in r.warnings]
    if all_warnings:
        _para(doc, s["warnings_head"] + ":", bold=True, color=_PLACEHOLDER_RED)
        for w in all_warnings:
            _bullet(doc, w, color=_PLACEHOLDER_RED)
    else:
        _para(doc, s["no_warnings"], italic=True, color=_GREY)
    # cross-references to detail sections (internal clickable links)
    see = "详见：" if lang == "zh" else "See also: "
    p = doc.add_paragraph()
    _style_run(p.add_run(see), size=10, bold=True, color=_NAVY)
    _add_internal_link(p, "sec3", s["sec3"], size=10)
    _style_run(p.add_run(" · "), size=10, color=_GREY)
    _add_internal_link(p, "sec4", s["sec4"], size=10)
    _style_run(p.add_run(" · "), size=10, color=_GREY)
    _add_internal_link(p, "sec5", s["sec5"], size=10)


def _add_overview(doc, model: RevenueModel, s: dict,
                  company_overview: Optional[str]):
    _heading(doc, s["sec2"], level=1, bookmark="sec2")
    if company_overview:
        _para(doc, company_overview)
    else:
        _para(doc, s["overview_placeholder"], italic=True,
              color=_PLACEHOLDER_RED)
    # segment share table (latest year)
    latest_year = model.years()[-1]
    total = sum(seg.revenue(latest_year) for seg in model.segments) or 1.0
    rows = []
    for seg in model.segments:
        rev = seg.revenue(latest_year)
        rows.append([seg.name, _yi(rev), f"{rev / total:.1%}"])
    _table(doc, [s["segment"], s["revenue_yi"], s["rev_share"]], rows)


def _add_driver_tables(doc, model: RevenueModel, s: dict, lang: str):
    _heading(doc, s["sec3"], level=1, bookmark="sec3")
    years = model.years()
    headers = [s["driver"], s["kind"], s["unit"], s["grade"]] + \
              [str(y) for y in years] + [s["source"]]
    ncols = len(headers)
    src_col = ncols - 1
    n_meta = 4  # driver/kind/unit/grade columns before the year columns
    for seg in model.segments:
        _heading(doc, seg.name, level=2)
        t = doc.add_table(rows=1, cols=ncols)
        t.style = "Table Grid"
        for ci, h in enumerate(headers):
            cell = t.rows[0].cells[ci]
            cell.text = ""
            _style_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=_NAVY)
        for d in seg.drivers():
            row = t.add_row()
            color = _LEVEL_COLOR.get(d.level, _COLOR_C)
            meta = [d.name, d.kind_label(lang), d.unit, d.level]
            for ci, val in enumerate(meta):
                cell = row.cells[ci]
                cell.text = ""
                _style_run(cell.paragraphs[0].add_run(str(val)), size=10, color=color)
            for yi, y in enumerate(years):
                cell = row.cells[n_meta + yi]
                cell.text = ""
                txt = (f"{d.values[y]:.4f}".rstrip("0").rstrip(".")
                       if y in d.values else "")
                _style_run(cell.paragraphs[0].add_run(txt), size=10, color=color)
            # source cell — clickable hyperlink when source_url is set
            src_cell = row.cells[src_col]
            src_cell.text = ""
            if d.source_url:
                _add_hyperlink(src_cell.paragraphs[0], d.source_url,
                               d.source or d.source_url, size=9)
            else:
                _style_run(src_cell.paragraphs[0].add_run(d.source or "—"),
                           size=10, color=color)
        # segment revenue row (computed)
        rev_row = t.add_row()
        rev_vals = [s["seg_revenue"], "", "", ""] + \
                   [_yi(seg.revenue(y)) for y in years] + [""]
        for ci, val in enumerate(rev_vals):
            cell = rev_row.cells[ci]
            cell.text = ""
            _style_run(cell.paragraphs[0].add_run(str(val)), size=10, bold=True)


def _add_residual(doc, model: RevenueModel, s: dict,
                  results: List[YearResult]):
    _heading(doc, s["sec4"], level=1, bookmark="sec4")
    rows = []
    colors = []
    for r in results:
        rows.append([
            str(r.year),
            _yi(r.segment_sum),
            _yi(r.total_revenue),
            _yi(r.residual),
            f"{r.residual_ratio:.1%}",
        ])
        colors.append(None)
    _table(doc, [s["year"], s["seg_sum"], s["total_reported"].split("(")[0].strip(),
                 s["residual"], s["residual_ratio"]], rows)
    for r in results:
        for w in r.warnings:
            _bullet(doc, f"{r.year}: {w}", color=_PLACEHOLDER_RED)


def _add_uncertainty(doc, model: RevenueModel, s: dict,
                     ranges: Optional[Dict[str, tuple]],
                     forecast_years: Optional[Sequence[int]],
                     include_charts: bool, lang: str):
    _heading(doc, s["sec5"], level=1, bookmark="sec5")
    latest_year = model.years()[-1]
    is_default = ranges is None
    eff_ranges = ranges if ranges is not None else _model_default_ranges(model, latest_year)
    # range-source callout (honest default annotation)
    _para(doc, s["ranges_default"] if is_default else s["ranges_user"],
          italic=True, color=_WARN_AMBER if is_default else _GREY)

    from .monte_carlo import simulate_model, tornado, scenarios
    mc = simulate_model(model, latest_year, eff_ranges, n=10000, seed=0)
    scs = scenarios(mc)

    # Bear/Base/Bull table
    _para(doc, s["scenarios_head"] + ":", bold=True)
    bb_rows = [[sc.name if lang == "en" else
                {"Bear": s["bear"], "Base": s["base"], "Bull": s["bull"]}[sc.name],
                _yi(sc.revenue), f"P{sc.percentile * 100:.0f}"] for sc in scs]
    _table(doc, ["", s["revenue_yi"], s["percentile"]], bb_rows)

    if include_charts:
        plt = _import_mpl()
        from .viz import (plot_revenue_distribution, plot_tornado, plot_forecast)
        # distribution
        _para(doc, s["dist_subtitle"] + ":", bold=True)
        fig = plt.figure()
        plot_revenue_distribution(mc, ax=fig.gca(), lang=lang)
        _embed_fig(doc, fig, plt)
        # tornado (first segment, latest year — the dashboard convention)
        if model.segments:
            seg_ranges = {k: v for k, v in eff_ranges.items()
                          if k in {d.name for d in model.segments[0].drivers()}}
            if seg_ranges:
                items = tornado(model.segments[0], latest_year, seg_ranges)
                _para(doc, s["tornado_subtitle"] + f" — {model.segments[0].name}:", bold=True)
                fig2 = plt.figure()
                plot_tornado(items, ax=fig2.gca(), lang=lang)
                _embed_fig(doc, fig2, plt)
        # forecast trajectory (case A/B/C handling)
        _para(doc, s["forecast_subtitle"] + ":", bold=True)
        if forecast_years is None:
            _para(doc, s["forecast_none"], italic=True, color=_GREY, size=9.5)
            fig3 = plt.figure()
            plot_forecast(model, ax=fig3.gca(), lang=lang)
            _embed_fig(doc, fig3, plt)
        elif _has_forecast_values(model, forecast_years):
            fig3 = plt.figure()
            plot_forecast(model, forecast_years=forecast_years, ax=fig3.gca(), lang=lang)
            _embed_fig(doc, fig3, plt)
        else:
            years_str = ", ".join(str(y) for y in forecast_years)
            _para(doc, s["forecast_placeholder"].format(years=years_str),
                  italic=True, color=_PLACEHOLDER_RED)
    else:
        _para(doc, s["lim_no_charts"], italic=True, color=_GREY)
        # percentile table as the chart-less fallback
        pct = mc.percentiles
        _table(doc, [s["percentile"], s["revenue_yi"]],
               [[k.upper(), _yi(v)] for k, v in pct.items()] +
               [["mean", _yi(mc.mean)], ["median", _yi(mc.median)], ["σ", _yi(mc.stdev)]])
        # forecast fallback
        if forecast_years is None:
            _para(doc, s["forecast_none"], italic=True, color=_GREY, size=9.5)
        elif not _has_forecast_values(model, forecast_years):
            years_str = ", ".join(str(y) for y in forecast_years)
            _para(doc, s["forecast_placeholder"].format(years=years_str),
                  italic=True, color=_PLACEHOLDER_RED)


def _add_limitations(doc, model: RevenueModel, s: dict,
                     is_default_ranges: bool, include_charts: bool,
                     results: List[YearResult]):
    _heading(doc, s["sec6"], level=1, bookmark="sec6")
    _para(doc, s["lim_intro"])
    # C-grade drivers
    c_drivers = []
    for seg in model.segments:
        for d in seg.drivers():
            if d.level == LEVEL_C:
                c_drivers.append(f"{d.name} [{seg.name}]")
    if c_drivers:
        _para(doc, s["lim_c_grade"], bold=True)
        for nm in c_drivers:
            _bullet(doc, nm, color=_COLOR_C)
    else:
        _bullet(doc, s["lim_none_c"])
    # ranges caveat
    _bullet(doc, s["lim_default_ranges"] if is_default_ranges else s["lim_user_ranges"])
    if not include_charts:
        _bullet(doc, s["lim_no_charts"])
    # residual caveat
    latest = results[-1]
    if latest.residual_ratio > 0.30:
        _bullet(doc, s["lim_residual_high"].format(ratio=latest.residual_ratio),
                color=_PLACEHOLDER_RED)
    else:
        _bullet(doc, s["lim_residual_ok"])
    _bullet(doc, s["lim_diagnostic_en"])


def _add_methodology(doc, s: dict, lang: str):
    _heading(doc, s["sec7"], level=1, bookmark="sec7")
    for line in s["method_body"]:
        _bullet(doc, line)
    # GitHub doc links — each principle one click from its full treatment
    ref = "详见项目文档：" if lang == "zh" else "See project docs: "
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    _style_run(p.add_run(ref), size=10, bold=True, color=_NAVY)
    _add_hyperlink(p, _DOC_URLS["design_principles"], "design-principles.md", size=10)
    _style_run(p.add_run(" · "), size=10, color=_GREY)
    _add_hyperlink(p, _DOC_URLS["industry_fit"], "industry-fit-analysis.md", size=10)
    _style_run(p.add_run(" · "), size=10, color=_GREY)
    _add_hyperlink(p, _DOC_URLS["extractor_proposal"], "proposal-segment-extraction.md", size=10)


def _add_source_index(doc, model: RevenueModel, s: dict, lang: str):
    """Data-source index appendix: every driver's source + URL, so each number
    in the memo is one click from its origin. Mirrors the revenue-model-builder
    skill's '链接：[URL]' convention."""
    title = "附录：数据来源索引" if lang == "zh" else "Appendix: Data Source Index"
    _heading(doc, title, level=1, bookmark="sec_sources")
    intro = ("下列为模型所有 driver 的数据来源，每个数字可溯源至原始出处（点击 URL 跳转）。"
             if lang == "zh" else
             "Every driver's data source; each number is traceable to its origin (click the URL).")
    _para(doc, intro, italic=True, color=_GREY, size=9.5)
    headers = [s["driver"], s["source"], "URL", s["grade"], s["segment"]]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        _style_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=_NAVY)
    seen = []
    for seg in model.segments:
        for d in seg.drivers():
            key = (d.source, d.source_url)
            if key not in seen:
                seen.append((d, seg.name))
    for d, segname in seen:
        row = t.add_row()
        color = _LEVEL_COLOR.get(d.level, _COLOR_C)
        # columns: 0=driver, 1=source, 2=URL, 3=grade, 4=segment
        plain = {0: d.name, 1: d.source or "—", 3: d.level, 4: segname}
        for ci, val in plain.items():
            cell = row.cells[ci]
            cell.text = ""
            _style_run(cell.paragraphs[0].add_run(str(val)), size=10, color=color)
        url_cell = row.cells[2]
        url_cell.text = ""
        if d.source_url:
            _add_hyperlink(url_cell.paragraphs[0], d.source_url, d.source_url, size=8.5)
        else:
            _style_run(url_cell.paragraphs[0].add_run("—"), size=10, color=_GREY)


# ---- public entry point -----------------------------------------------------
def build_docx(
    model: RevenueModel,
    path: str,
    *,
    forecast_years: Optional[Sequence[int]] = None,
    ranges: Optional[Dict[str, tuple]] = None,
    include_charts: bool = True,
    company_overview: Optional[str] = None,
    lang: Lang = "en",
) -> str:
    """Render ``model`` into a 7-section analyst memo at ``path``.

    Parameters mirror :func:`~revenue_model.excel_builder.build_excel` for zero
    learning cost. The memo covers: executive summary, company/segment
    overview, ABC-graded driver tables, residual alignment, uncertainty &
    scenarios (Monte Carlo + tornado + forecast), limitations, and a
    methodology appendix.

    Parameters
    ----------
    model : RevenueModel
        The model to render.
    path : str
        Output ``.docx`` path.
    forecast_years : sequence of int, optional
        Years treated as forecast. ``None`` → historical-only memo (Principle
        5). If passed but drivers lack values for those years, a prominent
        placeholder flags the gap instead of silently skipping.
    ranges : dict, optional
        ``driver_name -> (low, high)`` uncertainty intervals for Monte Carlo.
        ``None`` → default ±10% bands (illustrative only; flagged in §5 and §6).
        Pass real A/B/C-reflective intervals for analysis you'll defend.
    include_charts : bool, default True
        Embed matplotlib charts. ``False`` renders tables only (no matplotlib
        needed). If True but matplotlib is missing, raises ImportError with the
        install hint.
    company_overview : str, optional
        Narrative for §2. If None, leaves a TODO placeholder.
    lang : {"en", "zh"}, default "en"
        Memo language. ``"en"`` is the global default (PyPI/GitHub); pass
        ``"zh"`` for 中文版. The footnote on every memo shows the active
        language and how to switch.

    Returns
    -------
    str
        The output ``path`` (for chaining / CLI echo).
    """
    _check_lang(lang)
    s = _STRINGS[lang]
    results = model.validate_all()
    is_default_ranges = ranges is None

    doc = Document()
    # default font on the Normal style so unlabeled runs inherit YaHei too
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10.5)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), _FONT)
    rFonts.set(qn("w:ascii"), _FONT)
    rFonts.set(qn("w:hAnsi"), _FONT)

    _add_cover(doc, model, s, lang)
    _add_executive_summary(doc, model, s, lang, results)
    _add_overview(doc, model, s, company_overview)
    _add_driver_tables(doc, model, s, lang)
    _add_residual(doc, model, s, results)
    _add_uncertainty(doc, model, s, ranges, forecast_years, include_charts, lang)
    _add_limitations(doc, model, s, is_default_ranges, include_charts, results)
    _add_methodology(doc, s, lang)
    _add_source_index(doc, model, s, lang)

    doc.save(path)
    return path


def main(argv=None):
    """Module-level CLI: ``python -m revenue_model.docx_builder [output]``.

    Mirrors ``excel_builder.main``. The unified entry point
    (``python -m revenue_model docx ...``) lives in ``__main__.py``.
    """
    import argparse
    import os
    from .demo import build_novatech

    parser = argparse.ArgumentParser(
        description="Build the NovaTech demo revenue model to a Word memo (.docx)")
    parser.add_argument(
        "output", nargs="?", default="NovaTech_revenue_model_demo.docx",
        help="output .docx path (default: ./NovaTech_revenue_model_demo.docx)")
    parser.add_argument(
        "--lang", default="en", choices=["zh", "en"],
        help="memo language (default: en; zh = 中文版)")
    parser.add_argument(
        "--no-charts", action="store_true",
        help="render tables only, no embedded charts (skips matplotlib)")
    args = parser.parse_args(argv)

    model = build_novatech()
    out = os.path.abspath(args.output)
    build_docx(model, out, lang=args.lang, include_charts=not args.no_charts)
    print(f"OK -> {out}")
    print(f"language: {args.lang}  |  charts: {not args.no_charts}")


if __name__ == "__main__":
    main()
